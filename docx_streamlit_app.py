
import streamlit as st
from docx import Document
import pandas as pd
import re
import json
from typing import Dict, List, Tuple
from st_diff_viewer import diff_viewer
import difflib
from pathlib import Path
from HybridParser import HybridParser

def generate_medical_tree_html_universal(data_json, output_file="drzewo_programu.html"):
    if isinstance(data_json, str):
        try:
            data = json.loads(data_json)
        except:
            data = data_json
    else:
        data = data_json

    html = """<!DOCTYPE html>
<html lang="pl">
<head>
  <meta charset="UTF-8" />
  <title>🩺 Drzewo kwalifikacji – Program lekowy</title>
  <style>
    body { font-family: 'Segoe UI', sans-serif; margin: 30px; background: #f9f9f9; color: #333; }
    details { margin-left: 20px; border-left: 2px solid #007acc; padding-left: 15px; }
    summary { font-weight: bold; cursor: pointer; font-size: 1.1em; color: #007acc; margin-top: 10px; }
    ul { margin-top: 5px; margin-bottom: 10px; }
    li { margin: 4px 0; }
    .section-title { font-weight: bold; color: #2c3e50; }
  </style>
</head>
<body>
  <h1>🩺 Drzewo kwalifikacji – Program lekowy</h1>
  <p><strong>Instrukcja:</strong> Klikaj w strzałki, by rozwijać sekcje.</p>
"""

    section_titles = {
        "diagnosis": "🩺 Diagnoza",
        "age_requirements": "📅 Wymagania wiekowe",
        "eligibility_general": "✅ Ogólne kryteria kwalifikacji",
        "eligibility_detailed": "🔍 Szczegółowe kryteria",
        "treatment_rules": "📜 Zasady stosowania",
        "treatment_duration": "⏱️ Czas trwania leczenia",
        "discontinuation_criteria": "🛑 Kryteria wyłączenia"
    }

    for key, title in section_titles.items():
        if key in data and data[key]:
            items = data[key]
            if isinstance(items, list):
                html += f"""  <details open>
    <summary class="section-title">{title}</summary>
    <ul>
"""
                for item in items:
                    html += f"      <li>{item.strip(' .;')}</li>\n"
                html += "    </ul>\n  </details>\n"
            elif isinstance(items, str) and items.strip():
                html += f"""  <details open>
    <summary class="section-title">{title}</summary>
    <p>{items.strip()}</p>
  </details>\n"""
            elif isinstance(items, dict):
                html += f"""  <details open>
    <summary class="section-title">{title}</summary>
"""
                for subkey, subitems in items.items():
                    subname = "Chemioterapia i leczenie celowane" if "chemo" in subkey else "Immunoterapia"
                    html += f"    <details><summary>{subname}</summary><ul>\n"
                    for drug in subitems:
                        html += f"      <li>{drug}</li>\n"
                    html += "    </ul></details>\n"
                html += "  </details>\n"

    html += """
  <p style="margin-top: 30px; font-size: 0.9em; color: #666;">
    Wygenerowano automatycznie.
  </p>
</body>
</html>"""

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(html)

    print(f"✅ Wygenerowano: {output_file}")
    return html

def parse_medical_program_v2(text: str) -> dict:
    """
    Zaawansowany parser, który najpierw dzieli tekst na sekcje, potem je analizuje.
    Działa na RJG, hepatologię, itp.
    """
    # Normalizacja
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[.*?\]\(.*?\)', '', text)
    text = text.replace(';', '. ')  # zastąp średniki kropkami dla lepszego podziału

    # Podział na sekcje po nagłówkach
    sections = {}
    current_section = "start"
    sections[current_section] = []

    # Wzorce nagłówków
    patterns = [
        (r'I\.\s+w\s+zakresie', 'chemotherapy_targeted'),
        (r'II\.\s+w\s+zakresie', 'immunotherapy'),
        (r'W\s+leczeniu.*stosowane', 'treatment_rules'),
        (r'Kryteria\s+kwalifikacji', 'eligibility'),
        (r'Ogólne\s+kryteria', 'eligibility_general'),
        (r'Szczegółowe\s+kryteria', 'eligibility_detailed'),
        (r'Określenie\s+czasu\s+leczenia', 'treatment_duration'),
        (r'Kryteria\s+wyłączenia', 'discontinuation_criteria')
    ]

    lines = re.split(r'\.\s+(?=[A-Z])', text)  # podział po kropkach przed dużą literą
    full_text = ' '.join(lines)

    for pattern, key in patterns:
        match = re.search(pattern, full_text, re.I)
        if match:
            # Znajdź koniec sekcji (następny nagłówek lub koniec)
            next_patterns = [p for p, _ in patterns if p != pattern]
            next_match = None
            for p in next_patterns:
                m = re.search(p, full_text[match.end():], re.I)
                if m:
                    next_match = m
                    break
            end = match.end() + next_match.start() if next_match else len(full_text)
            sections[key] = full_text[match.end():end].strip()

    # Ekstrakcja leków
    def extract_drugs(text_chunk):
        # Szukaj po nazwach leków
        drugs = []
        known_drugs = ['aflibercept', 'pembrolizumab', 'niwolumab', 'ipilimumab', 
                      'triflurydyna', 'typiracyl', 'frukwintynib', 'bewacyzumab']
        for drug in known_drugs:
            if drug.lower() in text_chunk.lower():
                drugs.append(drug)
        return list(set(drugs))

    # Budowa JSON
    result = {
        "program_name": "Leczenie zaawansowanego raka jelita grubego",
        "diagnosis": ["Zaawansowany rak jelita grubego (IV stopień)"],
        "age_requirements": ["Wiek 18 lat i powyżej"],
        "eligibility_general": re.findall(r'•\s+(.*?)\.', sections.get('eligibility_general', '')) or 
                               [l.strip() for l in re.split(r';|\n', sections.get('eligibility_general', '')) if len(l.strip()) > 10],
        "eligibility_detailed": [sections.get('eligibility_detailed', '')],
        "financed_substances": {
            "chemotherapy_targeted": extract_drugs(sections.get('chemotherapy_targeted', '')),
            "immunotherapy": extract_drugs(sections.get('immunotherapy', ''))
        },
        "treatment_rules": sections.get('treatment_rules', ''),
        "treatment_duration": sections.get('treatment_duration', ''),
        "discontinuation_criteria": [l.strip() for l in re.split(r';|\n', sections.get('discontinuation_criteria', '')) if len(l.strip()) > 10]
    }

    return result

def parse_medical_guidelines(text: str) -> dict:
    """
    Parsuje tekst wytycznych medycznych, wyodrębniając kluczowe sekcje,
    takie jak finansowane leki, kryteria kwalifikacji i wyłączenia.

    Args:
        text (str): Surowy tekst z dokumentu.

    Returns:
        dict: Słownik zawierający ustrukturyzowane dane.
    """
    # Wstępne czyszczenie tekstu
    text = re.sub(r'\[.*?\]\(http.*?\)', '', text)  # Usunięcie linków markdown
    text = re.sub(r'(\n\s*){2,}', '\n\n', text)  # Normalizacja pustych linii

    # Struktura wyjściowa
    output = {
        "finansowane_substancje": {},
        "zasady_stosowania": {},
        "kryteria_kwalifikacji": {"ogolne": [], "szczegolowe": []},
        "czas_leczenia": "",
        "kryteria_wylaczenia": []
    }

    # Wyrażenia regularne do identyfikacji głównych sekcji
    patterns = {
        "chemioterapia": r'I\.\s+w\s+zakresie\s+chemioterapii\s+i\s+leczenia\s+celowanego:(.*?)(?=II\.)',
        "immunoterapia": r'II\.\s+w\s+zakresie\s+immunoterapii:(.*?)(?=W\s+leczeniu)',
        "zasady_stosowania": r'W\s+leczeniu\s+zaawansowanego\s+raka\s+jelita\s+grubego\s+stosowane\s+są:(.*?)(?=Kryteria\s+kwalifikacji)',
        "kryteria_ogolne": r'Ogólne\s+kryteria\s+kwalifikacji(.*?)(?=Szczegółowe\s+kryteria|Ponadto\s+do\s+programu)',
        "kryteria_szczegolowe": r'Szczegółowe\s+kryteria\s+kwalifikacji\s+do\s+terapii(.*?)(?=Określenie\s+czasu\s+leczenia)',
        "czas_leczenia": r'Określenie\s+czasu\s+leczenia(.*?)(?=Kryteria\s+wyłączenia)',
        "kryteria_wylaczenia": r'Kryteria\s+wyłączenia\s+z\s+programu(.*?)$'
    }

    # --- 1. Ekstrakcja list leków ---
    if match := re.search(patterns["chemioterapia"], text, re.S):
        drugs = [line.strip('; \n') for line in match.group(1).strip().split('\n') if line.strip()]
        output["finansowane_substancje"]["chemioterapia_i_leczenie_celowane"] = drugs
    
    if match := re.search(patterns["immunoterapia"], text, re.S):
        drugs = [line.strip('; \n.') for line in match.group(1).strip().split('\n') if line.strip()]
        output["finansowane_substancje"]["immunoterapia"] = drugs

    # --- 2. Ekstrakcja zasad stosowania ---
    if match := re.search(patterns["zasady_stosowania"], text, re.S):
        content = match.group(1).strip()
        # Lista leków kluczowych do podziału tekstu
        drug_names = ["pembrolizumab", "aflibercept", "niwolumab w skojarzeniu z ipilimumabem", "triflurydyna z typiracylem", "frukwintynib"]
        # Sortowanie od najdłuższego do najkrótszego, aby uniknąć błędnego dopasowania
        sorted_drug_names = sorted(drug_names, key=len, reverse=True)
        # Dynamiczne tworzenie wzorca do podziału
        split_pattern = f"({'|'.join(map(re.escape, sorted_drug_names))})"
        
        parts = re.split(split_pattern, content, flags=re.IGNORECASE)
        # Łączenie nazwy leku z jego opisem
        i = 1
        while i < len(parts) - 1:
            drug_name = parts[i].strip()
            description = parts[i+1].strip()
            output["zasady_stosowania"][drug_name.lower()] = description
            i += 2

    # --- 3. Ekstrakcja kryteriów ---
    def extract_list_items(content):
        # Dzieli tekst na punkty listy, ignorując puste linie
        return [item.strip('; ') for item in content.strip().split('\n') if item.strip() and len(item.strip()) > 3]

    if match := re.search(patterns["kryteria_ogolne"], text, re.S):
        output["kryteria_kwalifikacji"]["ogolne"] = extract_list_items(match.group(1))

    if match := re.search(patterns["kryteria_szczegolowe"], text, re.S):
        output["kryteria_kwalifikacji"]["szczegolowe"] = extract_list_items(match.group(1))
        
    if match := re.search(patterns["kryteria_wylaczenia"], text, re.S):
        output["kryteria_wylaczenia"] = extract_list_items(match.group(1))

    # --- 4. Czas leczenia ---
    if match := re.search(patterns["czas_leczenia"], text, re.S):
        output["czas_leczenia"] = match.group(1).strip()

    return output


def detect_headers(text: str) -> List[str]:
    """
    Wyszukuje potencjalne nagłówki:
    - linie pisane w całości wielkimi literami
    - linie kończące się dwukropkiem
    """
    lines = text.splitlines()
    headers = []
    for line in lines:
        t = line.strip()
        if not t: 
            continue
        if t.isupper() or t.endswith(':'):
            headers.append(re.escape(t))
    return headers

def split_by_headers(text: str) -> Dict[str, str]:
    """
    Dzieli tekst na sekcje według wykrytych nagłówków.
    Zwraca mapę nagłówek → zawartość.
    """
    headers = detect_headers(text)
    if not headers:
        return {'CAŁY_DOKUMENT': text.strip()}
    pattern = r'^(?:' + '|'.join(headers) + r')\s*$'
    parts = re.split(pattern, text, flags=re.MULTILINE)
    # findall to zachowanie kolejności nagłówków
    found = re.findall(pattern, text, flags=re.MULTILINE)
    sections = {}
    for idx, hdr in enumerate(found):
        body = parts[idx+1].strip()
        sections[hdr.strip()] = body
    # ostatnia część bez nagłówka na końcu
    tail = parts[-1].strip()
    if tail:
        sections.setdefault('BEZ_NAGŁÓWKA', tail)
    return sections

def extract_points(text: str) -> List[str]:
    """
    Ekstrakcja punktów listy:
    - po średnikach
    - punktach list ('-', '*', '+') na początku linii
    - numeracji '1.', 'a)', '1)'
    """
    # zastąp nowe linie spacjami
    norm = re.sub(r'\s*\n\s*', ' ', text)
    raw = re.split(r';\s*|\s*[-*+]\s*|\s*\d+\.\s*|\s*[a-z]\)\s*', norm)
    return [it.strip() for it in raw if it.strip()]

def extract_schedules(text: str) -> List[Tuple[str, List[str]]]:
    """
    Szuka fragmentów harmonogramu:
    - wzorzec czasu: 'w dniu', 'w X tygodniu', 'co X tygodni'
    - oddziela nagłówek czasu od listy punktów
    """
    # znajdź wszystkie propozycje kroków czasowych
    times = re.findall(r'(w dniu[^:;]+|w\s*\d+[, \d]*\s*tygodn[iu]|co\s*\d+\s*tygodn[iu])', text, flags=re.IGNORECASE)
    result = []
    for tm in times:
        # podział: nagłówek czasowy + reszta do następnego czasu
        parts = re.split(re.escape(tm), text, flags=re.IGNORECASE, maxsplit=1)
        if len(parts) < 2:
            continue
        after = parts[1]
        # listę kończymy przy kolejnym czasie lub końcu
        next_time = re.search(r'(w dniu[^:;]+|w\s*\d+[, \d]*\s*tygodn[iu]|co\s*\d+\s*tygodn[iu])', after, flags=re.IGNORECASE)
        segment = after if not next_time else after[:next_time.start()]
        items = extract_points(segment)
        result.append((tm.strip(), items))
    return result

def parse_document2(text: str) -> Dict[str, object]:
    """
    Zwraca strukturę:
    {
      'TYTUŁ_SEKCJI': ['punkt1','punkt2',…],
      'HARMONOGRAM': [('w dniu...', ['pt1', …]), ...],
      ...
    }
    """
    sections = split_by_headers(text)
    output: Dict[str, object] = {}
    for hdr, body in sections.items():
        pts = extract_points(body)
        if any(re.match(r'(w dniu|w\s*\d+|co\s*\d+)', s, re.IGNORECASE) for s in body.split()):
            # podejrzenie harmonogramu
            sched = extract_schedules(body)
            output[f'Harmonogram — {hdr}'] = sched
        else:
            output[hdr] = pts
    return output

def split_into_sections(text: str) -> Dranict[str, str]:
    """
    Dzieli tekst na sekcje, zakładając że nagłówki to całe linie wielkimi literami
    lub linie kończące się dwukropkiem.
    """
    # Znajdź wszystkie potencjalne nagłówki
    headers = re.findall(r'^(?:[A-ZĄĆĘŁŃÓŚŹŻ ]{5,}|.*?:)\s*$', text, flags=re.MULTILINE)
    
    # Jeśli brak nagłówków, całość jako jedna sekcja
    if not headers:
        return {'CAŁY_DOKUMENT': text.strip()}
    
    # Buduj mapę nagłówek → zawartość
    sections = {}
    # Stwórz wzorzec łączący kolejne nagłówki
    pattern = r'^(?P<header>' + '|'.join(map(re.escape, headers)) + r')\s*$'
    parts = re.split(pattern, text, flags=re.MULTILINE)
    # parts = [pusty, nagł1, zaw1, nagł2, zaw2, ..., ostatnia]
    for i in range(1, len(parts), 2):
        hdr = parts[i].strip()
        body = parts[i+1].strip()
        sections[hdr] = body
    return sections

def parse_section_items(section_text: str) -> List[str]:
    """
    Rozdziela akapit sekcji na listę punktów:
    - po średnikach
    - po wypunktowaniach (-, *) na początku linii
    - po numeracji (1., a))
    """
    # Usuń zbędne nowe linie w obrębie punktów
    normalized = re.sub(r'\s*\n\s*', ' ', section_text)
    # Podziel po średnikach lub punktach listy
    items = re.split(r';\s+|\n[-*]\s+|\d+\.\s+|[a-z]\)\s+', normalized)
    return [it.strip() for it in items if it.strip()]

def parse_document(text: str) -> Dict[str, List[str]]:
    """
    Zwraca słownik nagłówek → lista punktów z sekcji.
    """
    sections = split_into_sections(text)
    structured = {}
    for hdr, body in sections.items():
        items = parse_section_items(body)
        structured[hdr] = items
    return structured

def extract_title(doc):
    """Extract the title from the document"""
    # Look for the title pattern in the document
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "LECZENIE CHORYCH " in text and "ICD-10" in text:
            return text
    return None

def extract_table_data(doc):
    """Extract specific columns from tables"""
    extracted_data = {
        'ŚWIADCZENIOBIORCY': [],
        'SCHEMAT DAWKOWANIA LEKÓW W PROGRAMIE': [],
        'BADANIA DIAGNOSTYCZNE WYKONYWANE W RAMACH PROGRAMU': []
    }

    for table in doc.tables:
        # Check if this is the main table by looking for header patterns
        header_found = False
        for row in table.rows:
            for cell in row.cells:
                if 'ŚWIADCZENIOBIORCY' in cell.text:
                    header_found = True
                    break
            if header_found:
                break

        if header_found:
            # Extract data from the table
            rows = list(table.rows)
            if len(rows) > 1:
                # Find column indices
                header_row = rows[1]
                # if header_row.text=='ZAKRES ŚWIADCZENIA GWARANTOWANEGO' :
                #     header_row = rows[1] # Debugging line to see header row contents
                col_indices = {}

                for i, cell in enumerate(header_row.cells):
                    print(cell.text)  # Debugging line to see cell contents
                    cell_text = cell.text.strip()
                    if 'ŚWIADCZENIOBIORCY' in cell_text:
                        col_indices['ŚWIADCZENIOBIORCY'] = i
                    elif 'SCHEMAT DAWKOWANIA' in cell_text:
                        col_indices['SCHEMAT DAWKOWANIA LEKÓW W PROGRAMIE'] = i
                    elif 'BADANIA DIAGNOSTYCZNE' in cell_text:
                        col_indices['BADANIA DIAGNOSTYCZNE WYKONYWANE W RAMACH PROGRAMU'] = i

                # Extract data rows
                for row in rows[1:]:
                    for col_name, col_idx in col_indices.items():
                        if col_idx < len(row.cells):
                            extracted_data[col_name].append(row.cells[col_idx].text.strip())
                        else:
                            extracted_data[col_name].append("")

    return extracted_data

def process_docx_file(file_path):
    """Main function to process DOCX file"""
    try:
        doc = Document(file_path)

        # Extract title
        title = extract_title(doc)

        # Extract table data
        table_data = extract_table_data(doc)

        return title, table_data
    except Exception as e:
        return None, f"Error processing file: {str(e)}"

def main():
    st.set_page_config(page_title="Test czytanie wordów", page_icon="👻",layout="wide")
    st.title("Program Lekowy - Analiza Dokumentów DOCX")

    st.header("Wczytywanie pliku DOCX")

    docs_folder = Path("doks")

    if not docs_folder.exists():
        st.error(f"Folder 'doks/' nie istnieje.")
        return

    # Pobierz listę plików .docx
    docx_files = list(docs_folder.glob("*.docx"))

    # Selectbox do wyboru pliku
    selected_file_name = st.selectbox(
        "Wybierz plik DOCX z folderu 'doks/':",
        options=[f.name for f in docx_files],
        index=0
    )

    uploaded_file = None
    selected_file_path = docs_folder / selected_file_name
    uploaded_file = selected_file_path

    if uploaded_file is not None:
        # Save uploaded file temporarily
        # with open("temp_file.docx", "wb") as f:
        #     f.write(uploaded_file.getbuffer())

        # Process the file
        title, table_data = process_docx_file(selected_file_path)

        

        if title:
            st.header("Tytuł programu:")
            st.write(f"**{title}**")
        else:
            st.error("Nie można znaleźć tytułu programu w dokumencie.")

        if isinstance(table_data, dict) and any(table_data.values()):
            st.header("Dane z tabeli:")

            # Create DataFrame
            max_length = max(len(v) for v in table_data.values() if v)

            # Pad shorter lists with empty strings
            for key in table_data:
                while len(table_data[key]) < max_length:
                    table_data[key].append("")

            df = pd.DataFrame(table_data)

            test_HP = HybridParser()

            # Display as expandable sections
            #col1, col2, col3 = st.columns(3)
            col1 = st.container()
            col2 = st.container()
            col3 = st.container()

            with col1:
                st.subheader("ŚWIADCZENIOBIORCY")
                for i, content in enumerate(df['ŚWIADCZENIOBIORCY']):
                    if content.strip():
                        with st.expander(f"Sekcja {i+1}"):
                            #st.write(content)
                            data = parse_document(content)
                            string_wczytany = ""
                            for header, points in data.items():
                                st.write(f'## {header}')
                                string_wczytany += f'## {header}\n'
                                for idx, pt in enumerate(points, 1):
                                    st.write(f'{idx}. {pt}')
                                    string_wczytany += f'{idx}. {pt}\n'
                                st.write('---')
                                string_wczytany += '---\n'
                            
                            Htest1 = HybridParser.parse(test_HP,string_wczytany,1)

                            st.write(Htest1)

                            st.subheader("test:")
                            temp = parse_medical_program_v2(content)
                            st.json(temp)
                            html = generate_medical_tree_html_universal(json.dumps(temp))
                            with st.expander("Pokaż wygenerowany HTML"):
                                st.markdown(html, unsafe_allow_html=True)
                            if content and string_wczytany:
                                st.subheader("Różnice:")
                                # lines1 = content.splitlines()
                                # lines2 = string_wczytany.splitlines()
                                # if len(lines1) < len(lines2):
                                #     lines1.insert(0, "")
                                # elif len(lines2) < len(lines1):
                                #     lines2.insert(0, "")

                                # text1_aligned = "\n".join(lines1)
                                # text2_aligned = "\n".join(lines2)
                                # Użycie gotowego komponentu diff_viewer
                                diff_viewer(content, string_wczytany, split_view=True)

                                # Dodatkowe statystyki
                                similarity = difflib.SequenceMatcher(None, content, string_wczytany).ratio() 
                                st.metric("Podobieństwo", f"{similarity:.1%}")

            with col2:
                st.subheader("SCHEMAT DAWKOWANIA LEKÓW W PROGRAMIE")
                for i, content in enumerate(df['SCHEMAT DAWKOWANIA LEKÓW W PROGRAMIE']):
                    if content.strip():
                        with st.expander(f"Dawkowanie {i+1}"):
                            st.write(content)
                            if i==1:
                                Htest2 = HybridParser.parse(test_HP,content,2)

                                st.write(Htest2)

            with col3:
                st.subheader("BADANIA DIAGNOSTYCZNE WYKONYWANE W RAMACH PROGRAMU")
                for i, content in enumerate(df['BADANIA DIAGNOSTYCZNE WYKONYWANE W RAMACH PROGRAMU']):
                    if content.strip():
                        with st.expander(f"Badania {i+1}"):
                            st.write(content)
                            parsed = parse_document2(content)
                            for section, content in parsed.items():
                                st.write(f'## {section}')
                                if isinstance(content, list) and content and isinstance(content[0], tuple):
                                    # to harmonogram
                                    for time_label, items in content:
                                        st.write(f'- **{time_label}**:')
                                        for it in items:
                                            st.write(f'  - {it}')
                                else:
                                    for item in content if isinstance(content, list) else []:
                                        st.write(f'- {item}')
                                st.write('---')

                            if(i==1):
                                Htest3 = HybridParser.parse(test_HP,content,3)

                                st.write(Htest3)

            # Display full table
            st.header("Pełna tabela:")
            st.dataframe(df, use_container_width=True)

            # Download option
            csv = df.to_csv(index=False)
            st.download_button(
                label="Pobierz dane jako CSV",
                data=csv,
                file_name='program_lekowy_dane.csv',
                mime='text/csv'
            )
        else:
            st.warning("Nie znaleziono danych tabeli lub wystąpił błąd podczas ekstrakacji.")
            if isinstance(table_data, str):
                st.error(table_data)

if __name__ == "__main__":
    main()
